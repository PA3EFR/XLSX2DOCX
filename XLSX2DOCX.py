from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt


# Functie om een pagina-einde toe te voegen
def add_page_break(doc):
    doc.add_page_break()


# Excel bestand lezen
def read_xlsx(file_path):
    wb = load_workbook(file_path)
    sheet = wb.active
    rows = list(sheet.iter_rows(values_only=True))
    return rows


# Word-document maken
def create_word_from_xlsx(excel_file, word_file):
    rows = read_xlsx(excel_file)
    doc = Document()

    for row in rows:
        # Voeg rijgegevens toe als koptekst op de pagina
        header = doc.add_paragraph()
        header.alignment = 1  # Centreren
        run = header.add_run(" | ".join(str(cell) for cell in row if cell is not None))
        run.bold = True
        run.font.size = Pt(14)

        # Pagina-einde als er meer rijen zijn
        add_page_break(doc)

    # Opslaan van het Word-document
    doc.save(word_file)


# Bestanden instellen
excel_bestand = input("Geef het pad van het Excel-bestand op: ")
word_bestand = input("Geef de naam van het uitvoerbestand op (met .docx): ")

# Uitvoeren
create_word_from_xlsx(excel_bestand, word_bestand)
print(f"Word-document '{word_bestand}' is succesvol aangemaakt.")
